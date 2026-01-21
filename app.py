import streamlit as st
import pandas as pd
import random
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import io
import os

# ==================================================
# PAGE CONFIG
# ==================================================
st.set_page_config(page_title="Project Allocation", layout="centered")

# ==================================================
# HEADER
# ==================================================
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.image("vignan_logo.png", width=700)

st.markdown(
    "<h2 style='text-align:center;'> 📋 📑 🗂️ Project Allocation List</h2>",
    unsafe_allow_html=True
)
st.markdown("---")

uploaded = st.file_uploader("Upload Excel File (.xlsx)", type=["xlsx"])

# ==================================================
# CORRECTED BATCH LOGIC (STRICT PHASE ORDER)
# ==================================================
def generate_batches(df):
    df = df.copy()
    df["Category"] = df["Marks"].apply(lambda x: "Good" if x >= 7.0 else "Average")

    good = df[df["Category"] == "Good"].to_dict("records")
    avg = df[df["Category"] == "Average"].to_dict("records")

    random.shuffle(good)
    random.shuffle(avg)

    batches = []
    batch_no = 1

    # ==================================================
    # PHASE 1: GOOD + AVERAGE (ONLY 2 MEMBERS)
    # ==================================================
    while good and avg:
        batches.append([
            {"Batch": batch_no, **good.pop(0)},
            {"Batch": batch_no, **avg.pop(0)}
        ])
        batch_no += 1

    # ==================================================
    # PHASE 2: AVERAGE + AVERAGE (PAIRING)
    # ==================================================
    i = 0
    while i + 1 < len(avg):
        batches.append([
            {"Batch": batch_no, **avg[i]},
            {"Batch": batch_no, **avg[i + 1]}
        ])
        batch_no += 1
        i += 2

    leftover_avg = avg[i:]  # at most one

    # ==================================================
    # PHASE 3: FINAL ADJUSTMENT (ALLOW 3rd MEMBER)
    # ==================================================
    leftovers = leftover_avg + good

    for student in leftovers:
        placed = False
        for batch in batches:
            if len(batch) < 3:
                batch.append({"Batch": batch[0]["Batch"], **student})
                placed = True
                break
        if not placed:
            # safety fallback
            batches.append([
                {"Batch": batch_no, **student}
            ])
            batch_no += 1

    # ==================================================
    # ENSURE MINIMUM = 2
    # ==================================================
    final_batches = []
    single_students = []

    for batch in batches:
        if len(batch) == 1:
            single_students.extend(batch)
        else:
            final_batches.append(batch)

    for student in single_students:
        for batch in final_batches:
            if len(batch) < 3:
                batch.append({"Batch": batch[0]["Batch"], **student})
                break

    # ==================================================
    # OUTPUT FORMAT
    # ==================================================
    rows = []
    sno = 1
    for batch in final_batches:
        for s in batch:
            rows.append({
                "S.No": sno,
                "Batch No": s["Batch"],
                "Roll No": s["Roll No"],
                "Name": s["Name"]
            })
            sno += 1

    return pd.DataFrame(rows)

# ==================================================
# DOCX GENERATION
# ==================================================
def create_docx(df):
    doc = Document()
    if os.path.exists("vignan_logo.png"):
        doc.add_picture("vignan_logo.png", width=Inches(1))

    doc.add_heading("Project Allocation List", level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = col

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================================================
# PDF GENERATION
# ==================================================
def create_pdf(df):
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    if os.path.exists("vignan_logo.png"):
        elements.append(Image("vignan_logo.png", width=400, height=80))

    elements.append(Paragraph("<b>Project Allocation List</b>", styles["Heading2"]))
    elements.append(Paragraph("<br/>", styles["Normal"]))

    data = [df.columns.tolist()] + df.values.tolist()
    table = Table(data, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("GRID", (0,0), (-1,-1), 1, colors.black),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold")
    ]))

    elements.append(table)
    pdf.build(elements)
    buffer.seek(0)
    return buffer

# ==================================================
# PROCESS
# ==================================================
if uploaded:
    df = pd.read_excel(uploaded)

    if not {"Roll No", "Name", "Marks"}.issubset(df.columns):
        st.error("Excel must contain Roll No, Name, Marks (CGPA)")
        st.stop()

    allocation_df = generate_batches(df)

    st.success("✅ Batch Allocation Completed")
    st.dataframe(allocation_df)

    st.markdown("### Guide Allocation")
    guide1_end = st.number_input("Guide 1 – End Batch No", min_value=1, value=1)
    guide1_name = st.text_input("Guide Name (Batch 1 to End)")
    guide2_name = st.text_input("Guide Name (Remaining Batches)")

    if guide1_name and guide2_name:
        allocation_df["Guide Name"] = allocation_df["Batch No"].apply(
            lambda x: guide1_name if x <= guide1_end else guide2_name
        )

        st.success("✅ Guide Allocation Applied")
        st.dataframe(allocation_df)

        st.download_button("⬇️ Download DOCX", create_docx(allocation_df), "project_allocation.docx")
        st.download_button("⬇️ Download PDF", create_pdf(allocation_df), "project_allocation.pdf")

# ==================================================
# FOOTER
# ==================================================
st.markdown(
    """
    <style>
    .custom-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #0e2a3a;
        color: white;
        text-align: center;
        padding: 10px;
        font-size: 14px;
        z-index: 100;
    }
    </style>

    <div class="custom-footer">
        Developed by <b>Mr. A.N. Harshith Vardhan</b> | Department of Computer Applications
    </div>
    """,
    unsafe_allow_html=True
)
